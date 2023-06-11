# -*- coding: utf-8 -*-
import json
import re
import time

from docx import Document

import requests


def get_exam_id(subject_id, t):
    a = requests.get(
        'https://www.zlketang.com/wxpub/api/simulationv4?from=web&channel=web&devtype=web&platform_type=web&subject_id={}&t={}'.format(
            subject_id, t))
    exam_id_list = []
    if a.status_code == 200:
        response = json.loads(a.text)
        vip_data = response['data']['VIP保过题库']
        for per_data in vip_data:
            exam_id = per_data['exam_id']
            upgrade_info_str = per_data["upgrade_info"]
            upgrade_info_json = json.loads(upgrade_info_str)
            if upgrade_info_json['year'] == '2023':
                exam_id_list.append(exam_id)
    else:
        exam_id_list = '网络问题or反爬，采集失败！'
    return exam_id_list


def get_exam_detail(exam_id_list, cookie, referer):
    header = {
        "Accept": "*/*",
        "Accept-Encoding": "gzip, deflate, br",
        "Accept-Language": "zh-CN,zh-TW;q=0.9,zh;q=0.8",
        "Cookie": cookie,
        "Referer": referer,
        "Sec-Ch-Ua": "\"Google Chrome\";v=\"113\", \"Chromium\";v=\"113\", \"Not-A.Brand\";v=\"24\"",
        "Sec-Ch-Ua-Mobile": "?0",
        "Sec-Ch-Ua-Platform": "\"Windows\"",
        "Sec-Fetch-Dest": "empty",
        "Sec-Fetch-Mode": "cors",
        "Sec-Fetch-Site": "same-origin",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/113.0.0.0 Safari/537.36",
        "X-Requested-With": "XMLHttpRequest"
    }
    for exam_id in exam_id_list:
        print(exam_id)
        url = 'https://www.zlketang.com/wxpub/api/exam_question?exam_id={}&devtype=web'.format(exam_id)
        a = requests.get(url, headers=header)
        if a.status_code == 200:
            response = json.loads(a.text)
            try:
                exam_name = response['data']['exam_name']
                doc = Document()
                doc.add_paragraph('试卷名称：' + exam_name)
                print(exam_name)
                detail = response['data']['parts']
                for per_part in detail:
                    part_title = per_part['title']
                    part_description = per_part['description']
                    part_title_all = '{}（{}）'.format(part_title, part_description)
                    doc.add_paragraph(part_title_all)
                    if '单项选择题' in part_title or '判断题' in part_title:
                        doc.add_paragraph(f'<TYPE.TAG>单项选择题')
                    elif '多项选择题' in part_title:
                        doc.add_paragraph(f'<TYPE.TAG>多项选择题')
                    else:
                        doc.add_paragraph(f'<TYPE.TAG>文字题')
                    print(part_title_all)
                    questions = per_part['questions']
                    i = 1
                    for question in questions:
                        clean = re.compile('<.*?>')
                        question_title = question['description']
                        question_title = re.sub(clean, '', question_title)
                        question_title = f'{i}、' + question_title
                        doc.add_paragraph(question_title)
                        print(question_title)
                        options = question['options']
                        options = json.loads(re.sub(clean, '', options))
                        if '计算' in part_title or '综合题' in part_title or '简答题' in part_title:
                            j = 1
                            if type(options) == list:
                                for per_trouble in options:
                                    trouble_title = f'（{j}）' + per_trouble['description']
                                    j = j + 1
                                    print(trouble_title)
                                    doc.add_paragraph(trouble_title)
                            else:
                                doc.add_paragraph('')
                            answer = question['answer']
                            if '[' in answer:
                                answers = json.loads(answer)
                                doc.add_paragraph('参考答案:')
                                for answer in answers:
                                    doc.add_paragraph(answer)
                                    print(answer)
                            else:
                                doc.add_paragraph(answer)
                            solution = question['solution']
                            solution = re.sub(clean, '', solution)
                            if '[' in solution:
                                solutions = json.loads(solution)
                                doc.add_paragraph('解析:')
                                for solution in solutions:
                                    doc.add_paragraph(solution)
                                    print(solution)
                                doc.add_paragraph('分数:1')
                            else:
                                doc.add_paragraph(solution)
                                doc.add_paragraph('分数:1')
                            i = 1 + i
                        else:
                            for k, v in options.items():
                                option = '{}.{}'.format(k, v)
                                doc.add_paragraph(option)
                                print(option)
                            answer = question['answer'].replace(',', '')
                            answer = '参考答案:{}'.format(answer)
                            doc.add_paragraph(answer)
                            print(answer)
                            solution = question['solution']
                            solution = re.sub(clean, '', solution)
                            solution = '解析:{}'.format(solution)
                            doc.add_paragraph(solution)
                            doc.add_paragraph('分数:1')
                            print(solution)
                            i = 1 + i
                doc.save("{}.docx".format(exam_name))
                time.sleep(3)
            except Exception as e:
                print('报错如下：{}'.format(e))
        else:
            print('网络问题or反爬，采集失败！')


if __name__ == '__main__':
    subject_id = input('请输入参数subject_id：')
    t = input('请输入参数t：')
    a = get_exam_id(subject_id, t)
    print(a)
    if a == '网络问题or反爬，采集失败！':
        print('采集失败，已结束程序！！！')
        pass
    else:
        print('考卷id：{}'.format(a))
        cookie = input('请输入cookie：')
        referer = input('referer：')
        get_exam_detail(a, cookie, referer)
        print('===========程序结束！！==========')
