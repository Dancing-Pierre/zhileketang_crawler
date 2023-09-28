# -*- coding: utf-8 -*-
import json
import random
import re
import time

from docx import Document  # pip install python-docx

import requests


def get_exam_id(subject_id, t, exam_chinese_name):
    url = 'https://www.zlketang.com/wxpub/api/simulationv4?from=web&channel=web&devtype=web&platform_type=web&subject_id={}&t={}'.format(
        subject_id, t)
    a = requests.get(url=url, headers=header)
    exam_dict = {}
    get_name = False
    if a.status_code == 200:
        response = json.loads(a.text)
        vip_data = response['data']['VIP保过题库']
        for per_data in vip_data:
            exam_id = per_data['exam_id']
            exam_name = per_data['exam_name']
            upgrade_info_str = per_data["upgrade_info"]
            upgrade_info_json = json.loads(upgrade_info_str)
            if upgrade_info_json['year'] == '2023' or '押题' in exam_name or upgrade_info_json['year'] == 2023:
                if exam_name == exam_chinese_name or get_name:
                    get_name = True
                    exam_dict[f'{exam_name}'] = exam_id
    else:
        exam_dict = '网络问题or反爬，采集失败！'
    return exam_dict


def get_exam_detail(exam_dict):
    is_arrive = False
    for exam_name in exam_dict:
        exam_id = exam_dict[exam_name]
        print(f'开始采集试卷：{exam_id}')
        url = 'https://www.zlketang.com/wxpub/api/exam_question?exam_id={}&devtype=web'.format(exam_id)
        a = requests.get(url, headers=header)
        if a.status_code == 200:
            response = json.loads(a.text)
            exam_name = response['data']['exam_name']
            try:
                if not is_arrive:
                    if exam_chinese_name == exam_name:
                        is_arrive = True
                    else:
                        is_arrive = False
                if is_arrive:
                    doc = Document()
                    doc.add_paragraph('试卷名称：' + exam_name.replace('&nbsp;', ' '))
                    print(exam_name)
                    detail = response['data']['parts']
                    for per_part in detail:
                        part_title = per_part['title']
                        part_description = per_part['description']
                        part_title_all = '{}（{}）'.format(part_title, part_description)
                        doc.add_paragraph(part_title_all.replace('&nbsp;', ' '))
                        if '单项选择题' in part_title or '判断题' in part_title:
                            doc.add_paragraph(f'<TYPE.TAG>单项选择题')
                        elif '多项选择题' in part_title:
                            doc.add_paragraph(f'<TYPE.TAG>多项选择题')
                        else:
                            doc.add_paragraph(f'<TYPE.TAG>文字题')
                        print(part_title_all)
                        questions = per_part['questions']
                        i = 1
                        for question in questions:
                            clean = re.compile('<.*?>')
                            question_title = question['description']
                            question_title = re.sub(clean, '', question_title)
                            question_title = f'{i}、' + question_title
                            doc.add_paragraph(question_title.replace('&nbsp;', ' '))
                            print(question_title)
                            options = question['options']
                            options = json.loads(re.sub(clean, '', options))
                            if '计算' in part_title or '综合题' in part_title or '简答题' in part_title or '案例分析' in part_title:
                                j = 1
                                if type(options) == list:
                                    for per_trouble in options:
                                        trouble_title = f'（{j}）' + per_trouble['description']
                                        j = j + 1
                                        print(trouble_title)
                                        doc.add_paragraph(trouble_title.replace('&nbsp;', ' '))
                                else:
                                    doc.add_paragraph('')
                                answer = question['answer']
                                if '[' in answer:
                                    answers = json.loads(answer)
                                    doc.add_paragraph('参考答案:')
                                    for answer in answers:
                                        doc.add_paragraph(answer.replace('&nbsp;', ' '))
                                        print(answer)
                                else:
                                    doc.add_paragraph(answer.replace('&nbsp;', ' '))
                                solution = question['solution']
                                solution = re.sub(clean, '', solution)
                                if '[' in solution:
                                    solutions = json.loads(solution)
                                    doc.add_paragraph('解析:')
                                    for solution in solutions:
                                        doc.add_paragraph(solution.replace('&nbsp;', ' '))
                                        print(solution)
                                    doc.add_paragraph('分数:1')
                                else:
                                    doc.add_paragraph(solution.replace('&nbsp;', ' '))
                                    doc.add_paragraph('分数:1')
                                i = 1 + i
                            else:
                                for k, v in options.items():
                                    option = '{}.{}'.format(k, v)
                                    doc.add_paragraph(option.replace('&nbsp;', ' '))
                                    print(option)
                                answer = question['answer'].replace(',', '')
                                answer = '参考答案:{}'.format(answer)
                                doc.add_paragraph(answer.replace('&nbsp;', ' '))
                                print(answer)
                                solution = question['solution']
                                solution = re.sub(clean, '', solution)
                                solution = '解析:{}'.format(solution.replace('&nbsp;', ' '))
                                doc.add_paragraph(solution.replace('&nbsp;', ' '))
                                doc.add_paragraph('分数:1')
                                print(solution)
                                i = 1 + i
                    doc.save("{}.docx".format(exam_name))
                    time.sleep(random.randint(2, 4))
            except Exception as e:
                print('报错如下：{}，可以尝试换cookie重启程序，如不行再联系开发者！'.format(e))
                break
        else:
            print('网络问题or反爬，采集失败！')


if __name__ == '__main__':
    try:
        subject_id = input('请输入参数subject_id：')
        t = input('请输入参数t：')
        cookie = input('请输入cookie：')
        referer = input('请输入referer：')
        exam_chinese_name = input('请输入想要开始采集的试卷名：')
        header = {
            "Accept": "*/*",
            "Accept-Encoding": "gzip, deflate, br",
            "Accept-Language": "zh-CN,zh-TW;q=0.9,zh;q=0.8",
            "Cookie": cookie,
            "Referer": referer,
            "Sec-Ch-Ua": "\"Google Chrome\";v=\"113\", \"Chromium\";v=\"113\", \"Not-A.Brand\";v=\"24\"",
            "Sec-Ch-Ua-Mobile": "?0",
            "Sec-Ch-Ua-Platform": "\"Windows\"",
            "Sec-Fetch-Dest": "empty",
            "Sec-Fetch-Mode": "cors",
            "Sec-Fetch-Site": "same-origin",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/113.0.0.0 Safari/537.36",
            "X-Requested-With": "XMLHttpRequest"
        }
        a = get_exam_id(subject_id, t, exam_chinese_name)
        print(f'本次采集的试卷列表：{a}')
        if a == '网络问题or反爬，采集失败！':
            print('采集失败，已结束程序！！！')
            pass
        else:
            get_exam_detail(a)
            print('===========程序结束！！==========')
    except Exception as e:
        print('报错如下：{}，请联系开发者！'.format(e))
        time.sleep(30)
